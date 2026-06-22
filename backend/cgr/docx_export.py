"""DOCX 문서 생성 — 표준 임금명세서·근로계약서 다운로드용.

`python-docx` 활용. 텍스트 본문 → 양식 갖춘 .docx 바이트.

사용 패턴
    text = generate_service.run(...)   # 평문 본문 생성
    docx_bytes = text_to_docx(text, title='표준 임금명세서')
    # FastAPI 응답:
    return Response(docx_bytes, media_type=DOCX_MIMETYPE, headers={...})

설계
- 평문에서 `[섹션 제목]` 패턴을 헤딩으로 식별
- 한글 폰트 — '맑은 고딕' (Windows 기본) 적용
- A4 portrait, 표준 여백
"""
from __future__ import annotations

import io
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


DOCX_MIMETYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)

# 평문에서 헤딩 후보 패턴 — `[제목]` 또는 끝까지 `:` 인 짧은 줄
_HEADING_BRACKET = re.compile(r"^\s*\[([^\]]+)\]\s*$")
# `근 로 계 약 서`, `임 금 명 세 서` 같이 자간 띄운 제목 (가운데 정렬)
_HEADING_SPACED = re.compile(r"^\s*([가-힣]\s){2,}[가-힣]\s*$")


def _apply_korean_font(run, font_name: str = "맑은 고딕"):
    """run.font.name + east_asia rPr 둘 다 설정 — 한글 안 깨짐."""
    run.font.name = font_name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts"
    )
    if rFonts is None:
        from docx.oxml.ns import qn

        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.append(rFonts)
    from docx.oxml.ns import qn

    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)


def text_to_docx(
    body_text: str,
    *,
    title: str = "표준 문서",
    subtitle: str | None = None,
    footer_note: str | None = None,
) -> bytes:
    """평문 본문 → DOCX 바이트.

    Args:
        body_text: LLM 생성 본문.
        title: 문서 최상단 제목 (Heading 0).
        subtitle: 제목 하단 부제 (선택).
        footer_note: 본문 하단 안내 (선택).
    """
    doc = Document()

    # 페이지 설정 — A4 + 표준 여백
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    # ─── 제목 ───
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run(title)
    run_title.bold = True
    run_title.font.size = Pt(18)
    _apply_korean_font(run_title)

    if subtitle:
        p_sub = doc.add_paragraph()
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_sub = p_sub.add_run(subtitle)
        run_sub.font.size = Pt(10)
        run_sub.font.color.rgb = None  # 기본 회색 톤 — python-docx 기본 검정
        _apply_korean_font(run_sub)

    # 빈 줄
    doc.add_paragraph()

    # ─── 본문 ───
    for raw_line in body_text.splitlines():
        line = raw_line.rstrip()
        if not line:
            doc.add_paragraph()
            continue

        # 헤딩 — `[제목]`
        m1 = _HEADING_BRACKET.match(line)
        if m1:
            p = doc.add_paragraph()
            run = p.add_run(m1.group(1))
            run.bold = True
            run.font.size = Pt(13)
            _apply_korean_font(run)
            continue

        # 헤딩 — `근 로 계 약 서` (자간 띄운 제목, body 안 부속 제목)
        m2 = _HEADING_SPACED.match(line)
        if m2 and len(line.replace(" ", "")) <= 10:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line.strip())
            run.bold = True
            run.font.size = Pt(14)
            _apply_korean_font(run)
            continue

        # 본문 일반 — 한글 폰트 적용
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.size = Pt(10.5)
        _apply_korean_font(run)

    # ─── 푸터 안내 ───
    if footer_note:
        doc.add_paragraph()  # 공백
        p_foot = doc.add_paragraph()
        p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_foot = p_foot.add_run(footer_note)
        run_foot.italic = True
        run_foot.font.size = Pt(9)
        _apply_korean_font(run_foot)

    # 바이트로 직렬화
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ════════════════════════════════════════════════════════════════
# 취업규칙 신구대조표(3열 표) + 의견청취서 양식 DOCX
#   화면(WrComparisonView)과 동일한 표를 Word 표로 재현 → 표가 깨지지 않음.
#   다운로드 시 신구대조표 뒤에 '취업규칙 (개정) 의견청취서' 양식을 함께 첨부.
# ════════════════════════════════════════════════════════════════
def _fill_multiline(cell, text, *, size: float = 9.5, bold: bool = False, fill=None):
    """셀에 여러 줄(\n) 텍스트 채움 — 줄마다 단락 추가."""
    cell.text = ""
    lines = str(text or "").split("\n")
    for i, ln in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        run = p.add_run(ln)
        run.bold = bold
        run.font.size = Pt(size)
        _apply_korean_font(run)
    if fill:
        _shade(cell, fill)


def _append_opinion_form(doc) -> None:
    """취업규칙 (개정) 의견청취서 양식 (근로기준법 제94조)."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    SUB_FILL = "F3F4F6"

    pt = doc.add_paragraph()
    pt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rt = pt.add_run("취업규칙 (개정) 의견청취서")
    rt.bold = True
    rt.font.size = Pt(16)
    _apply_korean_font(rt)
    doc.add_paragraph()

    intro = doc.add_paragraph()
    ri = intro.add_run(
        "「근로기준법」 제94조에 따라 취업규칙의 작성·변경에 관하여 근로자 과반수"
        "(근로자 과반수로 조직된 노동조합이 있는 경우 그 노동조합)의 의견을 청취합니다."
    )
    ri.font.size = Pt(10)
    _apply_korean_font(ri)
    doc.add_paragraph()

    t = doc.add_table(rows=4, cols=2)
    t.style = "Table Grid"
    _fill(t.cell(0, 0), "사업장명", bold=True, center=True, fill=SUB_FILL)
    _fill(t.cell(0, 1), "")
    _fill(t.cell(1, 0), "대표자", bold=True, center=True, fill=SUB_FILL)
    _fill(t.cell(1, 1), "")
    _fill(t.cell(2, 0), "근로자 과반수(또는 노동조합) 대표", bold=True, center=True, fill=SUB_FILL)
    _fill(t.cell(2, 1), "")
    _fill(t.cell(3, 0), "의견청취일", bold=True, center=True, fill=SUB_FILL)
    _fill(t.cell(3, 1), "          년        월        일")

    doc.add_paragraph()
    pl = doc.add_paragraph()
    rl = pl.add_run("개정 취업규칙에 대한 의견")
    rl.bold = True
    rl.font.size = Pt(11)
    _apply_korean_font(rl)

    op = doc.add_table(rows=1, cols=1)
    op.style = "Table Grid"
    op.rows[0].cells[0].text = ""
    for _ in range(6):  # 의견 기재란 높이 확보
        op.rows[0].cells[0].add_paragraph()

    doc.add_paragraph()
    sign = doc.add_paragraph()
    sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rs = sign.add_run("근로자 과반수(또는 노동조합) 대표              (서명 또는 인)")
    rs.font.size = Pt(10.5)
    _apply_korean_font(rs)

    note = doc.add_paragraph()
    rn = note.add_run(
        "※ 취업규칙을 근로자에게 불리하게 변경하는 경우에는 의견청취가 아니라 "
        "근로자 과반수의 '동의'를 받아야 합니다(근로기준법 제94조 제1항 단서)."
    )
    rn.font.size = Pt(8.5)
    _apply_korean_font(rn)


def wr_comparison_to_docx(
    rows: "list[dict[str, Any]]",
    *,
    effective_date: str = "",
    footer_note: str | None = None,
) -> bytes:
    """취업규칙 신구대조표(3열 표) + 의견청취서 양식 → .docx 바이트."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK

    HEAD_FILL = "E5E7EB"
    rows = rows or []

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(1.6)
        section.right_margin = Cm(1.6)

    pt = doc.add_paragraph()
    pt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rt = pt.add_run("취업규칙 신구대조표")
    rt.bold = True
    rt.font.size = Pt(17)
    _apply_korean_font(rt)

    if (effective_date or "").strip():
        pdt = doc.add_paragraph()
        rdt = pdt.add_run(f"개정 취업규칙 시행일: {effective_date.strip()}")
        rdt.font.size = Pt(10)
        _apply_korean_font(rdt)
    doc.add_paragraph()

    n = len(rows)
    tbl = doc.add_table(rows=1 + max(n, 1), cols=3)
    tbl.style = "Table Grid"
    _fill(tbl.cell(0, 0), "개정 전 (현행)", bold=True, center=True, fill=HEAD_FILL)
    _fill(tbl.cell(0, 1), "개정 후 (개정안)", bold=True, center=True, fill=HEAD_FILL)
    _fill(tbl.cell(0, 2), "비고 (변경사유·관련 법령)", bold=True, center=True, fill=HEAD_FILL)

    if n == 0:
        c = tbl.cell(1, 0).merge(tbl.cell(1, 1)).merge(tbl.cell(1, 2))
        _fill(c, "변경이 필요한 조항이 없습니다.", center=True)
    else:
        for i, r in enumerate(rows):
            row = 1 + i
            head = " ".join(
                x for x in [str(r.get("article") or ""), str(r.get("title") or "")] if x
            ).strip()
            before = (head + "\n" if head else "") + str(r.get("before") or "")
            _fill_multiline(tbl.cell(row, 0), before)
            _fill_multiline(tbl.cell(row, 1), str(r.get("after") or ""))
            _fill_multiline(tbl.cell(row, 2), str(r.get("remark") or ""), size=9)

    widths = [Cm(7.0), Cm(7.0), Cm(4.6)]
    for row in tbl.rows:
        for idx, w in enumerate(widths):
            row.cells[idx].width = w

    # 페이지 나눔 → 의견청취서 양식
    pb = doc.add_paragraph()
    pb.add_run().add_break(WD_BREAK.PAGE)
    _append_opinion_form(doc)

    if footer_note:
        pf = doc.add_paragraph()
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rf = pf.add_run(footer_note)
        rf.italic = True
        rf.font.size = Pt(8.5)
        _apply_korean_font(rf)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ════════════════════════════════════════════════════════════════
# 공식 임금명세서 서식(표 레이아웃) DOCX — 구조화 form dict 를 칸별로 채움.
# 화면(WsPayslipFormView)과 동일한 레이아웃을 Word 표로 재현.
# ════════════════════════════════════════════════════════════════
from typing import Any  # noqa: E402


def _ws_amt(v: Any) -> str:
    s = str(v or "").strip()
    if not s:
        return ""
    digits = re.sub(r"[^\d.-]", "", s)
    if digits and re.fullmatch(r"-?\d+(\.\d+)?", digits):
        try:
            return f"{int(float(digits)):,}"
        except Exception:
            return s.replace("원", "")
    return s.replace("원", "")


def _ws_ym(form: dict[str, Any]) -> tuple[str, str]:
    src = str(form.get("settlementPeriod") or form.get("paymentDate") or "")
    m = re.search(r"(\d{4})\s*[-./년]\s*(\d{1,2})", src)
    if m:
        return m.group(1), str(int(m.group(2)))
    return "", ""


def _shade(cell, fill: str) -> None:
    from docx.oxml.ns import qn

    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.makeelement(qn("w:shd"), {})
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def _fill(cell, text, *, bold=False, size=10, center=False, right=False, fill=None):
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    cell.text = ""
    p = cell.paragraphs[0]
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif right:
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("" if text is None else str(text))
    run.bold = bold
    run.font.size = Pt(size)
    _apply_korean_font(run)
    if fill:
        _shade(cell, fill)


def payslip_form_to_docx(
    form: dict[str, Any],
    *,
    footer_note: str | None = None,
) -> bytes:
    """구조화 임금명세서 form → 공식 서식 표 레이아웃 .docx 바이트."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    HEAD_FILL = "E5E7EB"
    SUB_FILL = "F3F4F6"

    worker = form.get("worker") or {}
    emp = form.get("employer") or {}
    pays = form.get("payments") or []
    deds = form.get("deductions") or []
    y, mo = _ws_ym(form)

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    # 회사
    company = emp.get("company") or ""
    biz = emp.get("businessNo")
    comp_line = f"{company} ({biz})" if (company and biz) else company
    if comp_line:
        pc = doc.add_paragraph()
        rc = pc.add_run(comp_line)
        rc.font.size = Pt(10)
        _apply_korean_font(rc)

    # 제목
    pt = doc.add_paragraph()
    pt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rt = pt.add_run(f"{y or '____'} 년 {mo or '__'} 월 임금명세서")
    rt.bold = True
    rt.font.size = Pt(17)
    _apply_korean_font(rt)

    # 지급일
    pd = doc.add_paragraph()
    pd.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rd = pd.add_run(f"지급일 : {form.get('paymentDate') or ''}")
    rd.font.size = Pt(10)
    _apply_korean_font(rd)

    # 성명/사번/부서/직급
    info = doc.add_table(rows=2, cols=4)
    info.style = "Table Grid"
    _fill(info.cell(0, 0), "성명", bold=True, center=True, fill=SUB_FILL)
    _fill(info.cell(0, 1), worker.get("name"))
    _fill(info.cell(0, 2), "사번", bold=True, center=True, fill=SUB_FILL)
    _fill(info.cell(0, 3), worker.get("idOrBirth"))
    _fill(info.cell(1, 0), "부서", bold=True, center=True, fill=SUB_FILL)
    _fill(info.cell(1, 1), worker.get("dept"))
    _fill(info.cell(1, 2), "직급", bold=True, center=True, fill=SUB_FILL)
    _fill(info.cell(1, 3), worker.get("position"))

    doc.add_paragraph()

    # 세부 내역
    ph = doc.add_paragraph()
    ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rh = ph.add_run("세부 내역")
    rh.bold = True
    rh.font.size = Pt(11)
    _apply_korean_font(rh)

    n = max(len(pays), len(deds), 1)
    det = doc.add_table(rows=2 + n + 2, cols=4)
    det.style = "Table Grid"
    # 그룹 행
    g0 = det.cell(0, 0).merge(det.cell(0, 1))
    g1 = det.cell(0, 2).merge(det.cell(0, 3))
    _fill(g0, "지 급", bold=True, center=True, fill=SUB_FILL)
    _fill(g1, "공 제", bold=True, center=True, fill=SUB_FILL)
    # 컬럼 라벨
    _fill(det.cell(1, 0), "임금 항목", bold=True, center=True, fill=SUB_FILL)
    _fill(det.cell(1, 1), "지급 금액(원)", bold=True, center=True, fill=SUB_FILL)
    _fill(det.cell(1, 2), "공제 항목", bold=True, center=True, fill=SUB_FILL)
    _fill(det.cell(1, 3), "공제 금액(원)", bold=True, center=True, fill=SUB_FILL)
    # 데이터
    for i in range(n):
        r = 2 + i
        p = pays[i] if i < len(pays) else {}
        d = deds[i] if i < len(deds) else {}
        _fill(det.cell(r, 0), p.get("name"))
        _fill(det.cell(r, 1), _ws_amt(p.get("amount")), right=True)
        _fill(det.cell(r, 2), d.get("name"))
        _fill(det.cell(r, 3), _ws_amt(d.get("amount")), right=True)
    # 합계
    rt_ = 2 + n
    _fill(det.cell(rt_, 0), "지급액 계", bold=True, center=True, fill=SUB_FILL)
    _fill(det.cell(rt_, 1), _ws_amt(form.get("paymentTotal")), bold=True, right=True)
    _fill(det.cell(rt_, 2), "공제액 계", bold=True, center=True, fill=SUB_FILL)
    _fill(det.cell(rt_, 3), _ws_amt(form.get("deductionTotal")), bold=True, right=True)
    # 실지급액
    rn = 2 + n + 1
    blank = det.cell(rn, 0).merge(det.cell(rn, 1))
    _fill(blank, "")
    _fill(det.cell(rn, 2), "실지급액", bold=True, center=True, fill=HEAD_FILL)
    _fill(det.cell(rn, 3), _ws_amt(form.get("netPay")), bold=True, right=True, fill=HEAD_FILL)

    doc.add_paragraph()

    # 계산 방법
    pc2 = doc.add_paragraph()
    pc2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rc2 = pc2.add_run("계산 방법")
    rc2.bold = True
    rc2.font.size = Pt(11)
    _apply_korean_font(rc2)

    calc = [p for p in pays if str(p.get("basis") or "").strip()]
    ctbl = doc.add_table(rows=1 + max(len(calc), 1), cols=3)
    ctbl.style = "Table Grid"
    _fill(ctbl.cell(0, 0), "구분", bold=True, center=True, fill=SUB_FILL)
    _fill(ctbl.cell(0, 1), "산출식 또는 산출방법", bold=True, center=True, fill=SUB_FILL)
    _fill(ctbl.cell(0, 2), "지급액(원)", bold=True, center=True, fill=SUB_FILL)
    if calc:
        for i, p in enumerate(calc):
            _fill(ctbl.cell(1 + i, 0), p.get("name"))
            _fill(ctbl.cell(1 + i, 1), p.get("basis"))
            _fill(ctbl.cell(1 + i, 2), _ws_amt(p.get("amount")), right=True)
    else:
        c = ctbl.cell(1, 0).merge(ctbl.cell(1, 2))
        _fill(c, "별도 산출방법 기재 항목 없음", center=True)

    # 안내문
    pf = doc.add_paragraph()
    rf = pf.add_run(
        "※ 가족수당은 취업규칙 등에 지급요건이 규정되어 있는 경우 "
        "계산방법을 기재하지 않더라도 무방"
    )
    rf.font.size = Pt(8.5)
    _apply_korean_font(rf)

    if footer_note:
        pf2 = doc.add_paragraph()
        pf2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rf2 = pf2.add_run(footer_note)
        rf2.italic = True
        rf2.font.size = Pt(8.5)
        _apply_korean_font(rf2)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
